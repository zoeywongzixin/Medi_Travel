import os
import json
from pathlib import Path
from typing import Dict
from docx import Document
import ollama

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates" / "visa"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "generated_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_visa_document(user_data: Dict, package_data: Dict, origin_country: str) -> str:
    """
    Generates a fillable Word document using AI to naturally draft the letter 
    in the requested language based on the template guidelines.
    """
    template_path = TEMPLATES_DIR / f"{origin_country.lower()}.md"
    if not os.path.exists(template_path):
        template_path = TEMPLATES_DIR / "default.md"
        
    with open(template_path, "r", encoding="utf-8") as f:
        template_content = f.read()
        
    hospital_name = package_data.get("hospital", {}).get("name", "Unknown Hospital")
    city = package_data.get("hospital", {}).get("city", "Unknown City")
    
    # Use AI to generate the final text based on the template and parameters
    system_prompt = (
        "You are a professional administrative assistant for a Malaysian hospital. "
        "Your task is to draft a Medical Visa Invitation Letter. "
        "Use the provided template as a strict guideline for structure and language. "
        "Replace placeholders like [ENTER HOSPITAL NAME] and [ENTER CITY] with the provided real data. "
        "Leave user-specific placeholders (like [ENTER PATIENT NAME] or [ENTER PASSPORT NUMBER]) exactly as they are so the user can fill them in later. "
        "Do not include any external markdown formatting (like ```) in your final output, just raw text. "
        f"Ensure the language matches the intended language of the template ({origin_country} or English)."
    )
    
    user_prompt = f"""
    Hospital Data to inject:
    Hospital Name: {hospital_name}
    City: {city}
    
    Template:
    {template_content}
    """
    
    try:
        response = ollama.chat(
            model='llama3.2:3b',
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            options={'temperature': 0.1}
        )
        final_text = response['message']['content'].strip()
    except Exception as e:
        print(f"AI Generation failed, falling back to basic replace: {e}")
        final_text = template_content.replace("[ENTER HOSPITAL NAME]", hospital_name).replace("[ENTER CITY]", city)
        
    # Create Word Document
    doc = Document()
    
    # Split by newlines and add as paragraphs
    lines = final_text.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), 1)
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)
                
    filename = f"visa_application_{origin_country.lower()}.docx"
    output_path = OUTPUT_DIR / filename
    doc.save(str(output_path))
    
    return str(output_path)
